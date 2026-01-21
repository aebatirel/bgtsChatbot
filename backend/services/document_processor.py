"""Document processing service for parsing various file formats."""

import os
import uuid
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from backend.config import get_settings
from backend.utils.text_processing import clean_text, generate_preview


class DocumentProcessor:
    """Service for processing and extracting text from documents."""

    SUPPORTED_TYPES = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }

    def __init__(self):
        self.settings = get_settings()
        self._temp_files: dict[str, dict] = {}

    def is_supported(self, filename: str) -> bool:
        """Check if the file type is supported."""
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_TYPES

    def get_file_type(self, filename: str) -> str:
        """Get the MIME type for a file."""
        ext = Path(filename).suffix.lower()
        return self.SUPPORTED_TYPES.get(ext, "application/octet-stream")

    async def process_upload(
        self, file_content: bytes, filename: str
    ) -> dict:
        """Process an uploaded file and return extraction results."""
        ext = Path(filename).suffix.lower()

        # Extract text based on file type
        if ext == ".pdf":
            text = self._extract_pdf(file_content)
        elif ext in (".docx", ".doc"):
            text = await self._extract_docx(file_content, filename)
        elif ext in (".txt", ".md"):
            text = file_content.decode("utf-8", errors="ignore")
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        # Clean the extracted text
        text = clean_text(text)

        # Generate a unique upload ID
        upload_id = str(uuid.uuid4())

        # Store temporarily
        self._temp_files[upload_id] = {
            "filename": filename,
            "file_type": self.get_file_type(filename),
            "file_size": len(file_content),
            "full_text": text,
            "preview": generate_preview(text),
            "content": file_content,
        }

        return {
            "upload_id": upload_id,
            "filename": filename,
            "file_type": self.get_file_type(filename),
            "file_size": len(file_content),
            "preview": generate_preview(text),
            "full_text_length": len(text),
        }

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from a PDF file."""
        text_parts = []

        with fitz.open(stream=content, filetype="pdf") as doc:
            for page in doc:
                text_parts.append(page.get_text())

        return "\n\n".join(text_parts)

    async def _extract_docx(self, content: bytes, filename: str) -> str:
        """Extract text from a DOCX file."""
        # Save temporarily to read with python-docx
        temp_path = Path(self.settings.upload_dir) / f"temp_{uuid.uuid4()}.docx"

        try:
            temp_path.write_bytes(content)
            doc = DocxDocument(str(temp_path))

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)
        finally:
            if temp_path.exists():
                temp_path.unlink()

    def get_temp_file(self, upload_id: str) -> Optional[dict]:
        """Retrieve a temporarily stored file by upload ID."""
        return self._temp_files.get(upload_id)

    def remove_temp_file(self, upload_id: str) -> bool:
        """Remove a temporarily stored file."""
        if upload_id in self._temp_files:
            del self._temp_files[upload_id]
            return True
        return False

    def clear_temp_files(self):
        """Clear all temporary files."""
        self._temp_files.clear()


# Singleton instance
_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get the document processor singleton."""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
